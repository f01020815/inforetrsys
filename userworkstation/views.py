# Create your views here.
# -*- coding: utf-8 -*-
import os
import json
from django.shortcuts import render, render_to_response
from django.shortcuts import redirect
from adminworkstation.models import Path
from adminworkstation.models import PathApproval
from userworkstation.models import UploadFile
from workstation.publicwork.user_detail import user_detail
from DBoperation.DBuser.DBfeedbackwork import feedbackwork
from workstation.publicwork.html_to_path import html_to_path
from workstation.userwork.show_out import show_out
from DBoperation.DBuser.DBdelete_pathid import delete_pathid_work
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponse, Http404, FileResponse
from DBoperation.DBuser.DBuptdate_find_data import DBuptdate_find_data
from DBoperation.DBuser.DBuptdate_find_upload_file import DBuptdate_find_upload_file
from workstation.publicwork.log_record import normal, custom, normal_show
from userworkstation.models import Username
from DBoperation.DBuser.DBregister import insert_user
import docx
from DBoperation.DBadmin.DBapproval_find_add_del import DBapproval_find_add_del
from DBoperation.DBadmin.DBapproval_find_update import DBapproval_find_update
from workstation.publicwork.upload_high_file import upload_high_file
from django.contrib import messages
from workstation.userwork.update_get_new_list import update_get_new_list
from workstation.userwork.update_get_old_list import update_get_old_list
from workstation.userwork.update_nochange_upload import update_nochange_upload
from workstation.userwork.update_change_changefile import update_change_changefile
from workstation.publicwork.user_md5 import pwd_encrypt
from workstation.publicwork.user_admin_list import user_admin_list
from workstation.publicwork.user_list import user_list



# 意见反馈
def feedback(request):
    is_login = request.session.get('is_login', None)
    if is_login:
        if request.method == 'POST':
            # 获取用户输入的反馈内容 / 反馈人 / 反馈人联系方式 放入feedback_detail列表
            feedback_detail = [request.POST.get('feedback_words'),
                               request.POST.get('feedback_name'),
                               request.POST.get('feedback_phone')]

            feedback_detail.extend(user_detail(request))

            # insert FeedBack表
            feedbackwork(feedback_detail)

            normal(request, '意见增加', [request.POST.get('feedback_words'),
                               request.POST.get('feedback_name'),
                               request.POST.get('feedback_phone')], 0)
            results = "OK"
            return render(request, 'feedback.html', {'results': results})
        return render(request, 'feedback.html')
    else:
        return render(request, 'login.html', {'page': "feedback"})

# 普通用户单条数据插入
def insert(request):
    is_login = request.session.get('is_login', None)
    if is_login:
        inserted_path_id = 0
        #context = {}
        if request.method == 'POST':
            # 获取用户输入的连接类型 / 交易描述 / 交易类型 / 交易路径 / 备注说明 放入data_source列表
            data_source = [request.POST.get('line_type'),
                           request.POST.get('content'),
                           request.POST.get('trade_type'),
                           request.POST.get('trade_path'),
                           request.POST.get('remark')]
            source = 'html'

           # 获取上传高等级事件文件
            files = request.FILES.getlist("file123", None)

            # 检查上传的文件格式
            file_types = ['DOCX']
            for file in files:
                file_type_user = file.name.split('.')[len(file.name.split('.')) - 1]
                # 用户上传文件的扩展名和系统指定的扩展名进行'不区分大小写'的对比，对比一致则跳过
                if file_type_user.lower() in [file_type.lower() for file_type in file_types]:
                    pass
                else:
                    # 用户上传文件的扩展名和系统指定的扩展名进行'不区分大小写'的对比，对比不一致则返回文件名称
                    messages.success(request, "上传失败，上传文档格式错误，请上传DOCX格式文档!")
                    return render(request, 'insert.html')

            # insert Path表 SysPath表 PartWord表 PathApproval表
            inserted_path_id, new_approval_id = html_to_path(request, data_source, source)

            # 如果新增失败，返回值
            if inserted_path_id == 0:
                messages.success(request, "新增失败!可能已存在相同条目")
                return render(request, 'insert.html',
                              {'line_type': data_source[0], 'content': data_source[1], 'trade_type': data_source[2],
                              'trade_path': data_source[3], 'remark': data_source[4]})
            # 如果新增成功
            else:
                if len(files) == 0:
                    messages.success(request, "新增条目成功!")
                    return render(request, 'insert.html')
                else:
                    # 上传文档记库
                    approval_file_info = upload_high_file(inserted_path_id, files, request)
                    # 更新审批表approval_file_info
                    PathApproval.objects.filter(approval_id=new_approval_id).update(approval_n_file_info=str(approval_file_info))
                    messages.success(request, "新增条目成功!文档上传成功！")
        return render(request, 'insert.html')
    else:
        return render(request, 'login.html', {'page': "insert"})

# add by wangshibin 20190718
def index(request):

        return render(request, 'index.html')


# 查询结果展示
def show(request):

        datas = []
        key_words = []
        count = 0
        # index页面根据输入框获取的关键词进行查找，输出到show页面
        # 判断是否为POST方法传值
        if request.method == 'POST':
            # 获取输入框的内容
            inputs = request.POST.getlist('search-text')
            datas, key_words = show_out(inputs)
            if len(datas) == 0:
                #print("查询了零条记录")
                count = 0
                datas = "null"
            else:
                count = len(datas)
                # print("查询了%d条记录" % len(datas))
            # 查询记录日志
            normal_show(request, '查询', inputs, count)
        return render(request, 'show.html', {'shows': datas, 'count': count, 'get_names_words':  json.dumps(key_words, ensure_ascii=False)})


# add by wangshibin 20190722
@csrf_exempt
def delete_path_id(request):
    is_login = request.session.get('is_login', None)
    if is_login:

            if request.is_ajax():
                path_id = request.POST.get("path_id")
                # 判断是否为审批中
                res = Path.objects.filter(path_id=path_id)
                if res[0].path_approval_state == int(1):
                    log = []
                    for x in Path.objects.filter(path_id=path_id):
                        log.extend([x.path_id])
                        log.extend([x.path_linetype])
                        log.extend([x.path_content])
                        log.extend([x.path_tradetype])
                        log.extend([x.path_tradepath])
                        log.extend([x.path_remark])
                        approval_state = x.path_approval_state
                    log1 = log[1:5]
                    log2 = log[0]
                    normal(request, '删除记录', log1, log2)
                    # 获取用户信息
                    userinfo = user_detail(request)
                    # 更新Path表和审批表
                    delete_pathid_work(path_id, userinfo)
                    arg = 0
                else:
                    arg = 5
                return HttpResponse(arg)
    else:
        return render(request, 'login.html')

ajax_file_list=[]
@csrf_exempt
def update(request):
    is_login = request.session.get('is_login', None)
    if is_login:
        # 获取path_id值
        path_id = request.GET.get('path_id')
        # 获取文档变化后的file_id
        if request.is_ajax():
            global ajax_file_list
            ajax_file_list = request.POST.getlist("file_list")
        # 显示
        if request.method == 'GET':
            # 抓取修改前的path_id对应的各字段信息，显示在页面输入框上
            line_type, content, trade_type, trade_path, remark, state= DBuptdate_find_data(path_id)
            # 判断是否存在高等级事件报告
            get_reports_title_list = DBuptdate_find_upload_file(path_id)
            if get_reports_title_list:
                res_return = "havefile"
                return render(request, 'update.html', {'line_type': line_type, 'content': content, 'trade_type': trade_type,
                                                       'trade_path': trade_path, 'remark': remark, 'res_return': res_return,
                                                       'get_reports_title_list': get_reports_title_list,'state': state})
            else:
                res_return = "nofile"
                return render(request, 'update.html', {'line_type': line_type, 'content': content, 'trade_type': trade_type,
                                                       'trade_path': trade_path, 'remark': remark,
                                                       'res_return': res_return, 'state': state})

        # 更新
        elif request.method == 'POST':
            # 判断是否为审批中
            res = Path.objects.filter(path_id=path_id)
            if res[0].path_approval_state == int(0):
                messages.success(request, "条目审批中，无法修改！")
                new_address = ("/update/?path_id=" + str(path_id))
                return redirect(new_address)
            # 判断条目是否生效
            if len(res[0].path_hash) == 0:
                messages.success(request, "条目已删除，无法更新！")
                new_address = ("/update/?path_id=" + str(path_id))
                return redirect(new_address)




            # 获取上传的文件，如果没有文件，则默认为None
            upload_files = request.FILES.getlist("file", None)
            # 检查上传的文件格式,格式有误报错
            file_types = ['DOCX']
            for file in upload_files:
                file_type_user = file.name.split('.')[len(file.name.split('.')) - 1]
                # 用户上传文件的扩展名和系统指定的扩展名进行'不区分大小写'的对比，对比一致则跳过
                if file_type_user.lower() in [file_type.lower() for file_type in file_types]:
                    pass
                else:
                    # 用户上传文件的扩展名和系统指定的扩展名进行'不区分大小写'的对比，对比不一致则返回文件名称
                    messages.success(request, "上传失败，上传文档格式错误，请上传DOCX格式文档!")
                    new_address = ("/update/?path_id=" + str(path_id))
                    return redirect(new_address)

            # 获取输入框内容 连接类型 / 交易描述 / 交易类型 / 交易路径 / 备注说明 /hash / 结果fileid
            approval_n_list = update_get_new_list(request, ajax_file_list)

            # 获取原数据  连接类型 / 交易描述 / 交易类型 / 交易路径 / 备注说明 /hash / 数据库fileid
            approval_o_list = update_get_old_list(path_id)
            source = 'html'
            # 判断hash值和fileid是否一致
            # #####################################################################
            #    1：五个输入框未更新 & 未上传新文档   1.1   #
            #    2：五个输入框未更新 & 上传新文档     1.2   #
            #    3：五个输入框有更新 & 上传新文档     1.3   #
            #    4：五个输入框有更新 & 未上传新文档   1.4   #
            # #####################################################################
            # 1.1
            if approval_n_list[5] == approval_o_list[5] and approval_n_list[6] == approval_o_list[6] and len(upload_files)==0:
                messages.success(request, "更新失败!可能已存在该条目！")
                new_address = ("/update/?path_id=" + str(path_id))
                return redirect(new_address)
            # 1.2
            elif approval_n_list[5] == approval_o_list[5]:
                new_address = update_nochange_upload(path_id, approval_n_list, approval_o_list, request, upload_files)
                return redirect(new_address)
            # 1.3/1.4
            elif approval_n_list[5] != approval_o_list[5]:
                new_address = update_change_changefile(path_id, approval_n_list, approval_o_list, request, upload_files, source)
                return redirect(new_address)
        return render(request, 'update.html')
    else:
        return render(request, 'login.html', {'page': "update"})

# # 删除文件动作 数据文件夹里保留，只删除数据库  del by wangshibin 20190929
# def delete_file(request):
#     if request.method == 'GET':
#         file_id = request.GET.get('file_id')
#         new_url = DBdelete_fileid(request, file_id)
#         return redirect(new_url)
#     return HttpResponse()


# 下载文件动作
def download_file(request):
    file_id = request.GET.get('file_id')
    for x in UploadFile.objects.filter(file_id=int(file_id)):
        # for file in re.split(',', x.file_name):
        file_name = x.file_name
        file_paths = x.file_path
        file_path = os.path.join(file_paths, file_name)
        try:
            response = FileResponse(open(file_path, 'rb'))
            response['content_type'] = "application/octet-stream"
            response['Content-Disposition'] = 'attachment; filename = ' + os.path.basename(file_path).encode('utf-8').decode('ISO-8859-1')
            normal(request,'下载文件',file_name, 0)
            return response
        except Exception:
            raise Http404

# 回退条目展示
def rollback(request):
    is_login = request.session.get('is_login', None)
    if is_login:
        # 获取用户信息
        # 搜索审批表中的增加，删除项
        user = user_detail(request)
        approval_list_add_del = DBapproval_find_add_del(user)
        if len(approval_list_add_del) == 0:
            approval_list_add_del = ''
        # 搜索审批表中的更新项
        approval_list_update = DBapproval_find_update(user)
        if len(approval_list_update) == 0:
            approval_list_update = ''
        return render(request, 'rollback.html', {'approval_list_add_del': approval_list_add_del, 'approval_list_update': approval_list_update})
    else:
        return render(request, 'login.html', {'page': "rollback"})


def login(request):
    ret = {'status': ''}
    status = 'none'
    if request.method == 'POST':
        username = request.POST.get('username', None)
        password = request.POST.get('password', None)
        page = request.POST.get('page')
        if len(page) == 0 or page.isspace():
            page="login"
        page="/"+page+"/"
        password = pwd_encrypt(password)
        name = ''
        pwd = ''
        for x in Username.objects.filter(user_name=username):
            name = x.user_name
            pwd = x.user_password
            status = x.user_status
        if status is None:
            status = 'none'
        else:
            pass
        if status == 0:
            ret['status'] = 'account No approval'
        else:
            is_empty = all([username, password])
            if is_empty:
                if username == name:
                    pass
                    if password == pwd:
                        request.session['is_login'] = True
                        request.session['user_name'] = username
                        ret['status'] = 'login success'
                        return redirect(page)
                    else:
                        ret['status'] = 'password is error'
                else:
                    ret['status'] = 'account is error'
            else:
                ret['status'] = 'account or password is not empty'
    return render(request, "login.html", ret)

def register(request):
    is_login = request.session.get('is_login', None)
    if is_login:
        is_user = request.session.get('user_name', None)
        if is_user in user_admin_list():
            results = {}
            if request.method == 'POST':
                name = request.POST.get('user_name')
                pwd = request.POST.get('user_password')
                identity = request.POST.get('level')
                pwd = pwd_encrypt(pwd)
                level = request.POST.get('level')
                if name in user_list():
                    results['result'] = 'NO'
                    return render(request, 'register.html', results)
                else:
                    insert_user(name, pwd, level, '','', identity, 1)
                results['result'] = 'OK'
            return render(request, 'register.html', results)
        else:
            return render(request, 'login.html', {'page': "register"})
    else:
        return render(request, 'login.html', {'page': "register"})

def document_show(request):
    context = {}
    if request.method == 'GET':
        file_id = request.GET.get('file_id')
        for x in UploadFile.objects.filter(file_id=file_id):
            new_url = ("/update/?path_id=" + str(x.file_path_id))
        path = UploadFile.objects.filter(file_id=file_id)
        for i in path:
            file_name = i.file_name
            file_path = i.file_path
        doc = docx.Document(file_path+'\\'+file_name)

        ###文档标题###
        title = doc.paragraphs[0].text
        for i1 in doc.paragraphs:
            if i1.text == "一、事件基本情况":
                d1 = i1.text
            elif i1.text == "二、事件影响":
                d2 = i1.text
            elif i1.text == "三、事件损失评估":
                d3 = i1.text
            elif i1.text == "四、处理过程":
                d4 = i1.text
            elif i1.text == "五、事件处置分析":
                d5 = i1.text
            elif i1.text == "六、后续措施和建议":
                d6 = i1.text
        # for i2 in doc.paragraphs:
        #     if i2.text == "二、事件影响":
        #         d2 = i2.text
        #         print(i2.text)
        #
        # for i3 in doc.paragraphs:
        #     if i3.text == "三、事件损失评估":
        #         d3 = i3.text
        #         print(i3.text)
        #
        # for i4 in doc.paragraphs:
        #     if i4.text == "四、处理过程":
        #         d4 = i4.text
        #         print(i4.text)
        #
        #
        # for i5 in doc.paragraphs:
        #     if i5.text == "五、事件处置分析":
        #         d5 = i5.text
        #         print(i5.text)
        #
        # for i6 in doc.paragraphs:
        #     if i6.text == "六、后续措施和建议":
        #         d6 = i6.text
        #         print(i6.text)

        ###文档标题end###

        ###文档正文###
        for x in range(len(doc.paragraphs)):
            if doc.paragraphs[x].text == "一、事件基本情况":
                e1 = x
            elif doc.paragraphs[x].text == "二、事件影响":
                e2 = x
            elif doc.paragraphs[x].text == "三、事件损失评估":
                e3 = x
            elif doc.paragraphs[x].text == "四、处理过程":
                e4 = x
            elif doc.paragraphs[x].text == "五、事件处置分析":
                e5 = x
            elif doc.paragraphs[x].text == "六、后续措施和建议":
                e6 = x
        text1t = []
        text2t = []
        text3t = []
        text4t = []
        text5t = []
        text6t = []
        for i in range(e1 + 1, e2):
            text1t.append(doc.paragraphs[i].text)
            text1 = "\n".join(text1t)
        for i in range(e2 + 1, e3):
            text2t.append(doc.paragraphs[i].text)
            text2 = "\n".join(text2t)
        for i in range(e3 + 1, e4):
            text3t.append(doc.paragraphs[i].text)
            text3 = "\n".join(text3t)
        for i in range(e4 + 1, e5):
            text4t.append(doc.paragraphs[i].text)
            text4 = "\n".join(text4t)
        for i in range(e5 + 1, e6):
            text5t.append(doc.paragraphs[i].text)
            text5 = "\n".join(text5t)
        for i in range(e6 + 1, len(doc.paragraphs)):
            text6t.append(doc.paragraphs[i].text)
            text6 = "\n".join(text6t)
        context = {'title':title,
                   'k1':text1, 'k2':text2, 'k3':text3, 'k4':text4, 'k5':text5, 'k6':text6,
                   'data1':d1, 'data2':d2, 'data3':d3, 'data4':d4, 'data5':d5, 'data6':d6}
        return render(request, 'document_show.html',  context)
        # return redirect(new_url, context)
        # return render_to_response("update.html", context)


def register_user(request):
    if request.method == 'POST':
        user_name = request.POST.get('registerUsername')
        user_password = request.POST.get('registerPassword')
        user_password = pwd_encrypt(user_password)
        user_phone = request.POST.get('registerPhone')
        user_email = request.POST.get('registerEmail')
        user_department = request.POST.get('registerDepartment')
        user_identity = '普通用户'
        user_status = 0
        if user_name in user_list():
            messages.success(request, "用户重复")
            return render(request, 'register_user.html')
        else:
            messages.success(request, "注册成功")
        insert_user(user_name, user_password, user_phone, user_email, user_department, user_identity, user_status)
    return render(request, 'register_user.html')


