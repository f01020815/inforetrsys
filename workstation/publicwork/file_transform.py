import docx
from docx.shared import Inches
import re
from userworkstation.models import UploadFile


# doc = get_docx("E:\中行项目_python\\2019年6月5日POS交易异常事件处理分析报告V0.1.docx")
# for i in doc.parag

file_id = 1
path = UploadFile.objects.filter(file_id=2)
str = ""
arr = []

for i1 in doc.paragraphs:
    if i1.text == "一、事件基本情况":
        str = ""
        d1 = i1.text
    elif i1.text == "二、事件影响":
        d1 = str;
        str = "";
    elif i1.text == "三、事件损失评估":
        d2 = str;
        str = "";
    elif i1.text == "四、处理过程":
        d3 = str;
        str = "";
    elif i1.text == "五、事件处置分析":
        d4 = str;
        str = "";
    else:
        str+= i1.text;

d5 = str;
# for z in zz:
#     qq = z + "\n"
# data = []
# for i in doc.paragraphs:
#     if i.text == "二、事件影响":
#
#         break
#
#     print(i)
#     print(type(i))
# print(d1)

#
# for i in doc.paragraphs:
#     if i.text == "四、处理过程":
#         print(i.text)
#
# for i in doc.paragraphs:
#     if i.text == "五、事件处置分析":
#         print(i.text)
#
# for i in doc.paragraphs:
#     if i.text == "六、后续措施和建议":
#         print(i.text)
#
# print(data)
